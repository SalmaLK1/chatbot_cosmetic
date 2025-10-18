import os
import hashlib
import tempfile
from docx import Document as DocxDocument
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# ==============================
# FONCTIONS UTILITAIRES FICHIERS
# ==============================

def get_file_hash(file_bytes):
    """
    Génère un hash MD5 unique basé sur le contenu du fichier.
    Utilisé pour l'identification et la déduplication des documents.
    
    Args:
        file_bytes (bytes): Contenu brut du fichier
    
    Returns:
        str: Hash MD5 hexadécimal du fichier
    """
    return hashlib.md5(file_bytes).hexdigest()

def get_title_from_filename(filename):
    """
    Extrait un titre lisible depuis le nom de fichier.
    Supprime l'extension et les chemins de répertoire.
    
    Args:
        filename (str): Nom complet du fichier avec chemin
    
    Returns:
        str: Titre extrait sans extension ni chemin
    """
    return os.path.splitext(os.path.basename(filename))[0]

# ==============================
# GÉNÉRATION DE FICHIERS D'EXPORT
# ==============================

def generate_export_file(data, format="txt"):
    """
    Génère un fichier d'export dans différents formats à partir des données RAG.
    Supporte TXT, DOCX et PDF avec mise en forme adaptée.
    
    Args:
        data (dict): Dictionnaire contenant:
            - answer (str): Réponse générée par le modèle
            - context (list): Liste des documents de contexte utilisés
        format (str): Format de sortie ("txt", "docx", "pdf")
    
    Returns:
        str: Chemin vers le fichier temporaire généré
    
    Raises:
        ValueError: Si le format demandé n'est pas supporté
    """
    # Extraction des données avec valeurs par défaut
    answer = data.get("answer", "Aucune réponse")
    context = data.get("context", [])

    # ==============================
    # FORMAT TEXTE (TXT)
    # ==============================
    if format == "txt":
        """
        Format texte simple - lisible et universel
        Structure linéaire avec séparateurs clairs
        """
        temp = tempfile.NamedTemporaryFile(
            delete=False, 
            suffix=".txt", 
            mode="w", 
            encoding="utf-8"  # Support des caractères spéciaux
        )
        
        # Écriture de la réponse principale
        temp.write("Réponse générée :\n")
        temp.write(answer + "\n\n")
        
        # Ajout du contexte si disponible
        if context:
            temp.write("Contexte utilisé :\n")
            for i, doc in enumerate(context):
                # Numérotation et séparation des chunks
                temp.write(f"Chunk {i+1} : {doc.page_content}\n---\n")
        
        temp.close()
        return temp.name

    # ==============================
    # FORMAT DOCUMENT WORD (DOCX)
    # ==============================
    elif format == "docx":
        """
        Format DOCX avec mise en forme structurée
        Hiérarchie des titres et paragraphes
        """
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        
        # Création du document Word
        doc = DocxDocument()
        
        # Titre principal
        doc.add_heading("Réponse générée", 0)  # Niveau 0 = titre principal
        
        # Réponse principale
        doc.add_paragraph(answer)
        
        # Section contexte si disponible
        if context:
            doc.add_heading("Contexte utilisé", level=1)  # Niveau 1 = section
            
            for i, doc_chunk in enumerate(context):
                # Sous-titre pour chaque chunk
                doc.add_paragraph(f"Chunk {i+1} :", style='Heading2')
                # Contenu du chunk
                doc.add_paragraph(doc_chunk.page_content)
        
        # Sauvegarde du document
        doc.save(temp.name)
        return temp.name

    # ==============================
    # FORMAT PDF
    # ==============================
    elif format == "pdf":
        """
        Format PDF avec mise en forme professionnelle
        Utilise ReportLab pour la génération
        """
        temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        
        # Configuration du document PDF
        doc = SimpleDocTemplate(temp.name)
        
        # Récupération des styles par défaut
        styles = getSampleStyleSheet()
        
        # Construction des éléments du document
        flowables = [
            # Titre principal
            Paragraph("Réponse générée", styles["Heading1"]),
            # Réponse principale
            Paragraph(answer, styles["Normal"]),
        ]
        
        # Ajout de la section contexte si disponible
        if context:
            flowables.append(Paragraph("Contexte utilisé", styles["Heading2"]))
            
            for i, doc_chunk in enumerate(context):
                # Sous-titre pour chaque chunk
                flowables.append(Paragraph(f"Chunk {i+1} :", styles["Heading3"]))
                # Contenu du chunk
                flowables.append(Paragraph(doc_chunk.page_content, styles["Normal"]))
        
        # Génération du PDF
        doc.build(flowables)
        return temp.name

    # ==============================
    # GESTION DES FORMATS NON SUPPORTÉS
    # ==============================
    else:
        raise ValueError("Format non pris en charge")